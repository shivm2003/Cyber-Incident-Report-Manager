import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import tempfile
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_excel_report(incidents):
    """Generates a professional Excel report including AI analysis and 'Happened At' timestamps."""
    data = [{
        "Event Date": i.happened_at.strftime("%Y-%m-%d %H:%M") if i.happened_at else "Unknown",
        "Sync Date": i.date_collected.strftime("%Y-%m-%d %H:%M"),
        "Title": i.title,
        "Severity": i.severity or "Low",
        "Attack Type": i.attack_type or "Unknown",
        "Country": i.country,
        "Financial Sector": i.financial_sector if i.is_financial else "N/A",
        "AI Impact Summary": i.impact_summary or "Analyzing...",
        "Source": i.source,
        "Link": i.link
    } for i in incidents]
    
    df = pd.DataFrame(data)
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "cyber_intelligence_report.xlsx")
    df.to_excel(file_path, index=False)
    return file_path


def _draw_report_header_footer(canvas, doc):
    """Helper to draw header/footer on every page."""
    canvas.saveState()
    width, height = letter
    
    # Footer
    canvas.setFont("Helvetica-Bold", 8)
    canvas.setFillColor(colors.grey)
    canvas.drawCentredString(width / 2, 30, f"Page {canvas.getPageNumber()} | Confidential Security intelligence | Official Report by Shivam Mishra")
    
    # Header Accent
    canvas.setFillColor(colors.HexColor("#4f46e5"))
    canvas.rect(0, height - 2, width, 2, fill=1)
    
    # Subtle Watermark
    canvas.setFont("Helvetica-Bold", 60)
    canvas.saveState()
    canvas.translate(width/2, height/2)
    canvas.rotate(45)
    canvas.setFillColor(colors.grey, alpha=0.03)
    canvas.drawCentredString(0, 0, "SHIVAM MISHRA")
    canvas.restoreState()
    
    canvas.restoreState()

def generate_pdf_report(incidents, report_title="Cyber Intelligence Report"):
    """Generates a professional, high-fidelity PDF report with full intelligence details."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "cyber_intelligence.pdf")
    
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=letter,
        rightMargin=50, leftMargin=50,
        topMargin=70, bottomMargin=50
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'MainTitle',
        parent=styles['Heading1'],
        fontSize=26,
        textColor=colors.HexColor("#0a0e17"),
        alignment=TA_LEFT,
        spaceAfter=10
    )
    
    subtitle_style = ParagraphStyle(
        'SubTitle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        spaceAfter=30
    )
    
    incident_title_style = ParagraphStyle(
        'IncidentTitle',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=colors.HexColor("#1e1b4b"),
        spaceBefore=15,
        spaceAfter=5
    )
    
    meta_style = ParagraphStyle(
        'MetaStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
        fontName='Helvetica-Oblique'
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceBefore=8,
        spaceAfter=10
    )

    severity_styles = {
        "CRITICAL": ParagraphStyle('Crit', parent=body_style, textColor=colors.red, fontName='Helvetica-Bold'),
        "HIGH": ParagraphStyle('High', parent=body_style, textColor=colors.orange, fontName='Helvetica-Bold'),
        "MEDIUM": ParagraphStyle('Med', parent=body_style, textColor=colors.HexColor("#b45309"), fontName='Helvetica-Bold'),
        "LOW": ParagraphStyle('Low', parent=body_style, textColor=colors.HexColor("#059669"), fontName='Helvetica-Bold'),
    }

    elements = []
    
    # Report Header
    elements.append(Paragraph(report_title.upper(), title_style))
    elements.append(Paragraph(f"GLOBAL THREAT INTELLIGENCE FEED | GENERATED: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}", subtitle_style))
    elements.append(Spacer(1, 10))

    for idx, i in enumerate(incidents):
        severity = (i.severity or "Low").upper()
        sev_style = severity_styles.get(severity, severity_styles["LOW"])
        
        # Incident Title
        elements.append(Paragraph(f"#{idx+1} {i.title}", incident_title_style))
        
        # Metadata Row
        happened_str = i.happened_at.strftime("%Y-%m-%d %H:%M") if i.happened_at else "Recent"
        meta_text = f"<b>STATUS:</b> <font color='{sev_style.textColor}'>{severity}</font> | <b>DATE:</b> {happened_str} | <b>SOURCE:</b> {i.source} | <b>TARGET:</b> {i.country}"
        elements.append(Paragraph(meta_text, meta_style))
        
        # Intelligence Summary (NOT TRUNCATED)
        summary_text = i.ai_summary or i.impact_summary or i.description or "No detailed analysis available."
        elements.append(Paragraph(summary_text, body_style))
        
        # Separator line
        elements.append(Spacer(1, 5))
        line_table = Table([['']], colWidths=[letter[0]-100], rowHeights=[1])
        line_table.setStyle(TableStyle([('LINEBELOW', (0,0), (-1,-1), 0.5, colors.lightgrey)]))
        elements.append(line_table)
        elements.append(Spacer(1, 10))

    doc.build(elements, onFirstPage=_draw_report_header_footer, onLaterPages=_draw_report_header_footer)
    return file_path


def _draw_impact_report_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(colors.HexColor("#3b82f6")) # Blue
    canvas.rect(0, letter[1] - 4, letter[0], 4, fill=1, stroke=0)
    
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor("#94a3b8"))
    footer_text = "Generated by Shivam AI | Confidential Security Intelligence Impact Report"
    canvas.drawString(40, 30, footer_text)
    canvas.line(40, 45, letter[0]-40, 45)
    
    page_num = canvas.getPageNumber()
    canvas.drawRightString(letter[0]-40, 30, f"Page {page_num}")
    canvas.restoreState()

def generate_single_impact_pdf(report, mitre_mappings, forensic_content=None):
    """Generates a detailed, premium PDF for a single AI Impact Analysis report."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"AI_Impact_Report_{report.id}.pdf")
    
    from reportlab.platypus import HRFlowable

    doc = SimpleDocTemplate(
        file_path, 
        pagesize=letter,
        rightMargin=40, leftMargin=40,
        topMargin=50, bottomMargin=60
    )
    styles = getSampleStyleSheet()
    
    # Custom Modern Styles
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor("#0f172a"), spaceAfter=2, fontName="Helvetica-Bold", leading=28)
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor("#64748b"), spaceAfter=15, fontName="Helvetica")
    
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor("#1e293b"), spaceBefore=18, spaceAfter=10, fontName="Helvetica-Bold", textTransform='uppercase')
    h3_style = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=12, textColor=colors.HexColor("#334155"), spaceBefore=12, spaceAfter=6, fontName="Helvetica-Bold")
    
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10.5, leading=16, textColor=colors.HexColor("#334155"))
    body_bold = ParagraphStyle('BodyBold', parent=body_style, fontName="Helvetica-Bold")
    
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor("#64748b"), textTransform='uppercase', spaceAfter=2, fontName="Helvetica-Bold")
    
    mono_style = ParagraphStyle('Mono', parent=styles['Normal'], fontSize=9, fontName="Courier", textColor=colors.HexColor("#0f172a"), wordWrap='CJK')
    link_style = ParagraphStyle('Link', parent=styles['Normal'], fontSize=9.5, fontName="Courier", textColor=colors.HexColor("#3b82f6"), wordWrap='CJK')
    
    elements = []
    
    incident = getattr(report, 'incident', None)
    sev = (incident.severity or "LOW").upper() if incident else "LOW"
    
    sev_bg = "#f59e0b"
    if sev == "CRITICAL": sev_bg = "#ef4444"
    elif sev == "HIGH": sev_bg = "#f97316"
    elif sev == "MEDIUM": sev_bg = "#eab308"
    elif sev == "LOW": sev_bg = "#22c55e"
    
    # Header Section
    elements.append(Paragraph("AI IMPACT ANALYSIS", label_style))
    title_p = Paragraph(report.incident_title or "Cyber Impact Forensic Deep-Dive", title_style)
    sub_p = Paragraph(f"<b>REPORT #{report.id}</b> &bull; INCIDENT #{report.incident_id} &bull; PUBLISHED", subtitle_style)
    
    badge_data = [[
        Paragraph(f"<font size=10 color='white'><b>SEVERITY</b></font><br/><br/><font size=14 color='white'><b>{sev}</b></font>", ParagraphStyle('b', alignment=TA_CENTER))
    ]]
    badge_table = Table(badge_data, colWidths=[80], rowHeights=[70])
    badge_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,0), colors.HexColor(sev_bg)),
        ('ALIGN', (0,0), (0,0), 'CENTER'),
        ('VALIGN', (0,0), (0,0), 'MIDDLE'),
        ('ROUNDEDCORNERS', [8, 8, 8, 8]),
    ]))

    header_layout = Table([[ [title_p, sub_p], badge_table ]], colWidths=[400, 100])
    header_layout.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (1,0), (1,0), 'RIGHT'),
    ]))
    elements.append(header_layout)
    
    # Custom Divider
    elements.append(Spacer(1, 10))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=0, spaceAfter=20))
    
    # Intelligence Summary
    if incident and (incident.ai_summary or incident.description):
        elements.append(Paragraph("INTELLIGENCE SUMMARY", h2_style))
        summary_text = incident.ai_summary or incident.description
        elements.append(Paragraph(summary_text, body_style))
        elements.append(Spacer(1, 20))
        
    # AI Forensic Deep-Dive Analysis
    if forensic_content:
        forensic_header = ParagraphStyle('ForensicHeader', parent=h2_style, textColor=colors.HexColor("#7c3aed"))
        elements.append(Paragraph("AI FORENSIC ANALYSIS (DEEP-DIVE)", forensic_header))
        for paragraph in forensic_content.split('\n'):
            paragraph = paragraph.strip()
            if paragraph:
                elements.append(Paragraph(paragraph, body_style))
        elements.append(Spacer(1, 20))

    # Grid info
    inc_date = incident.happened_at.strftime("%Y-%m-%d %H:%M") if incident and incident.happened_at else "Unknown"
    cap_date = incident.date_collected.strftime("%Y-%m-%d %H:%M") if incident else "Unknown"
    entity = incident.target_entity or incident.country if incident else "Unknown"
    source = incident.source or 'Unknown' if incident else "Unknown"
    
    grid_data = [
        [Paragraph("INCIDENT DATE", label_style), Paragraph("CAPTURED DATE", label_style)],
        [Paragraph(f"<b>{inc_date}</b>", body_style), Paragraph(f"<b>{cap_date}</b>", body_style)],
        [Paragraph("TARGET ENTITY", label_style), Paragraph("SOURCE", label_style)],
        [Paragraph(f"<b>{entity}</b>", body_style), Paragraph(f"<b>{source}</b>", body_style)]
    ]
    grid_table = Table(grid_data, colWidths=[250, 250])
    grid_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
        ('LINEABOVE', (0,2), (-1,2), 1, colors.HexColor("#e2e8f0")),
        ('LINEBEFORE', (1,0), (1,-1), 1, colors.HexColor("#e2e8f0")),
        ('PADDING', (0, 0), (-1, -1), 12),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,0), 2),
        ('BOTTOMPADDING', (0,2), (-1,2), 2),
    ]))
    elements.append(grid_table)
    elements.append(Spacer(1, 20))

    # Source Evidence
    if incident and incident.link:
        elements.append(Paragraph("SOURCE EVIDENCE", h2_style))
        elements.append(Paragraph(f"&bull; <a href='{incident.link}'>{incident.link}</a>", link_style))
        elements.append(Spacer(1, 20))

    # Executive Summary
    elements.append(Paragraph("EXECUTIVE SUMMARY", h2_style))
    elements.append(Paragraph(report.official_report or "No official report generated.", body_style))
    elements.append(Spacer(1, 20))

    # Business Impact Analysis Grid
    elements.append(Paragraph("BUSINESS IMPACT ANALYSIS", h2_style))
    impact_data = [
        [Paragraph("BUSINESS IMPACT", label_style), Paragraph("OPERATIONAL IMPACT", label_style)],
        [Paragraph(f"{report.business_impact or 'N/A'}", body_style), Paragraph(f"{report.operational_impact or 'N/A'}", body_style)],
        [Paragraph("FINANCIAL IMPACT", label_style), Paragraph("REPUTATIONAL IMPACT", label_style)],
        [Paragraph(f"{report.financial_impact or 'N/A'}", body_style), Paragraph(f"{report.reputational_impact or 'N/A'}", body_style)]
    ]
    it = Table(impact_data, colWidths=[250, 250])
    it.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
        ('LINEABOVE', (0,2), (-1,2), 1, colors.HexColor("#e2e8f0")),
        ('LINEBEFORE', (1,0), (1,-1), 1, colors.HexColor("#e2e8f0")),
        ('PADDING', (0, 0), (-1, -1), 12),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('BOTTOMPADDING', (0,2), (-1,2), 6),
    ]))
    elements.append(it)
    elements.append(Spacer(1, 25))

    # Vectors and Data Exposure
    vectors_data = [
        [Paragraph("ROOT CAUSE", label_style), Paragraph(report.root_cause or 'Unknown', body_style)],
        [Paragraph("ATTACK TYPE", label_style), Paragraph(report.attack_type or 'Unknown', body_style)],
        [Paragraph("BREACH METHOD", label_style), Paragraph(report.breach_method or 'Unknown', body_style)],
        [Paragraph("DATA INVOLVED", label_style), Paragraph(report.data_involved or 'Unknown', body_style)],
        [Paragraph("CLASSIFICATION", label_style), Paragraph(report.data_classification or 'Unknown', body_style)],
        [Paragraph("AFFECTED ENTITIES", label_style), Paragraph(report.affected_customers or 'Unknown', body_style)],
    ]
    vt = Table(vectors_data, colWidths=[150, 350])
    vt.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
        ('PADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(Paragraph("TECHNICAL VECTORS & DATA EXPOSURE", h2_style))
    elements.append(vt)
    elements.append(Spacer(1, 25))

    # Deep-Dive Analysis
    elements.append(Paragraph("DEEP-DIVE ANALYSIS", h2_style))
    elements.append(Paragraph(report.technical_analysis or "No technical analysis available.", body_style))
    elements.append(Spacer(1, 25))

    # MITRE Mapping
    if mitre_mappings:
        elements.append(Paragraph("MITRE ATT&CK INTELLIGENCE MAPPING", h2_style))
        mitre_data = []
        for m in mitre_mappings:
            mitre_data.append([
                Paragraph(f"<b>{m.tactic}</b><br/><font size=8 color='#64748b'>{m.technique_id}</font>", body_style),
                Paragraph(f"<b>{m.technique_name}</b><br/>{m.analysis_justification}", body_style)
            ])
        if mitre_data:
            mt = Table(mitre_data, colWidths=[150, 350])
            mt.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            elements.append(mt)
        elements.append(Spacer(1, 20))

    # Breach Process Timeline
    elements.append(Paragraph("BREACH TIMELINE & PROCESS", h2_style))
    process_data = report.breach_process
    steps = []
    if process_data:
        import json
        try:
            parsed = json.loads(process_data)
            if isinstance(parsed, list):
                steps = parsed
            elif isinstance(parsed, dict):
                steps = list(parsed.values())
        except Exception:
            if isinstance(process_data, str):
                steps = [s.strip() for s in process_data.split('\n') if s.strip()]
    
    if steps:
        step_data = [[Paragraph(f"<b>{idx}</b>", body_style), Paragraph(step, body_style)] for idx, step in enumerate(steps, 1)]
        st = Table(step_data, colWidths=[30, 470])
        st.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('PADDING', (0,0), (-1,-1), 4)
        ]))
        elements.append(st)
    else:
        elements.append(Paragraph("No timeline available.", body_style))
    elements.append(Spacer(1, 20))

    # Crawled Intelligence Source Content
    if incident and incident.crawled_content:
        elements.append(Paragraph("CRAWLED INTELLIGENCE SOURCE CONTENT", h2_style))
        for p_text in incident.crawled_content.split('\n'):
            p_text = p_text.strip()
            if p_text:
                elements.append(Paragraph(p_text, ParagraphStyle('CrawledStyle', parent=body_style, fontSize=9, textColor=colors.HexColor("#64748b"))))
                elements.append(Spacer(1, 4))
        elements.append(Spacer(1, 20))

    doc.build(elements, onFirstPage=_draw_impact_report_footer, onLaterPages=_draw_impact_report_footer)
    return file_path


def generate_single_impact_docx(report, mitre_mappings, forensic_content=None):
    """Generates a professional Word document for a single AI Impact Analysis report."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"AI_Impact_Report_{report.id}.docx")
    
    doc = Document()
    
    # Title
    doc.add_heading('CYBER IMPACT FORENSIC REPORT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(report.incident_title or "Incident Analysis Deep-Dive", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run(f"REPORT ID: #{report.id} | INCIDENT ID: #{report.incident_id}\n").bold = True
    p.add_run(f"CONFIDENTIAL SECURITY INTELLIGENCE\n")
    p.add_run(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Metadata Table
    incident = getattr(report, 'incident', None)
    doc.add_heading('INCIDENT METADATA', level=2)
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    def fill_cell(r, c, label, value):
        cell = table.cell(r, c)
        cell.text = label
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        
        cell = table.cell(r, c+1)
        cell.text = str(value or "N/A")
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    if incident:
        source_val = f"{incident.source or 'Unknown'} (Link: {incident.link})" if incident.link else (incident.source or "Unknown")
        fill_cell(0, 0, "Source:", source_val)
        fill_cell(0, 2, "Severity:", (incident.severity or "Low").upper())
        fill_cell(1, 0, "Entity:", incident.target_entity or incident.country)
        fill_cell(1, 2, "Incident Date:", incident.happened_at.strftime("%Y-%m-%d %H:%M") if incident.happened_at else "Recent")
        fill_cell(2, 0, "Captured Date:", incident.date_collected.strftime("%Y-%m-%d %H:%M"))
        fill_cell(2, 2, "Breach Method:", report.breach_method)

    doc.add_paragraph()
    
    # Intelligence Summary & Evidence
    if incident:
        doc.add_heading('INTELLIGENCE SUMMARY', level=2)
        doc.add_paragraph(incident.ai_summary or incident.description or "N/A")
        
        if incident.link:
            doc.add_heading('SOURCE EVIDENCE', level=2)
            p = doc.add_paragraph()
            p.add_run("Intelligence Link: ").bold = True
            p.add_run(incident.link)
    
    # AI Forensic Deep-Dive Analysis (optional)
    if forensic_content:
        doc.add_heading('AI FORENSIC ANALYSIS (DEEP-DIVE)', level=2)
        for paragraph in forensic_content.split('\n'):
            paragraph = paragraph.strip()
            if paragraph:
                p = doc.add_paragraph(paragraph)
                for run in p.runs:
                    run.font.size = Pt(10)
        doc.add_paragraph()  # spacer

    # Executive Summary
    doc.add_heading('EXECUTIVE SUMMARY', level=2)
    doc.add_paragraph(report.official_report or "No official report generated.")
    
    # Impact Analysis
    doc.add_heading('BUSINESS IMPACT ANALYSIS', level=2)
    impacts = [
        ("Business Impact", report.business_impact),
        ("Operational Impact", report.operational_impact),
        ("Financial Impact", report.financial_impact),
        ("Reputational Impact", report.reputational_impact)
    ]
    for label, val in impacts:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(val or "N/A"))
        
    # Technical Vectors
    doc.add_heading('TECHNICAL VECTORS', level=2)
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run(str(report.root_cause or "Unknown"))
    
    p = doc.add_paragraph()
    p.add_run("Attack Type: ").bold = True
    p.add_run(str(report.attack_type or "Unknown"))
    
    p = doc.add_paragraph()
    p.add_run("Breach Method: ").bold = True
    p.add_run(str(report.breach_method or "Unknown"))
    
    # Data Exposure
    doc.add_heading('DATA EXPOSURE', level=2)
    p = doc.add_paragraph()
    p.add_run("Data Involved: ").bold = True
    p.add_run(str(report.data_involved or "Unknown"))
    
    p = doc.add_paragraph()
    p.add_run("Classification: ").bold = True
    p.add_run(str(report.data_classification or "Unknown"))
    
    p = doc.add_paragraph()
    p.add_run("Affected Entities: ").bold = True
    p.add_run(str(report.affected_customers or "Unknown"))
    
    doc.add_heading('DEEP-DIVE ANALYSIS', level=2)
    doc.add_paragraph(report.technical_analysis or "No technical analysis available.")
    
    # MITRE Mapping
    if mitre_mappings:
        doc.add_heading('MITRE ATT&CK INTELLIGENCE MAPPING', level=2)
        for m in mitre_mappings:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{m.tactic}: ").bold = True
            p.add_run(f"{m.technique_id} - {m.technique_name}")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(f"Justification: {m.analysis_justification}")
            run.italic = True
            run.font.size = Pt(9)
            
    # Breach Process
    doc.add_heading('BREACH TIMELINE & PROCESS', level=2)
    process_data = report.breach_process
    steps = []
    if process_data:
        import json
        try:
            parsed = json.loads(process_data)
            if isinstance(parsed, list):
                steps = parsed
            elif isinstance(parsed, dict):
                steps = list(parsed.values())
        except Exception:
            if isinstance(process_data, str):
                steps = [s.strip() for s in process_data.split('\n') if s.strip()]
    if steps:
        for idx, step in enumerate(steps, 1):
            doc.add_paragraph(f"{idx}. {step}")
    else:
        doc.add_paragraph("No timeline available.")
    
    # Executive Statement
    doc.add_heading('EXECUTIVE STATEMENT', level=2)
    p = doc.add_paragraph()
    p.add_run(f'"{report.official_report}"').italic = True

    # Crawled Intelligence Source Content
    if incident and incident.crawled_content:
        doc.add_heading('CRAWLED INTELLIGENCE SOURCE CONTENT', level=2)
        for line in incident.crawled_content.split('\n'):
            line = line.strip()
            if line:
                p_crawled = doc.add_paragraph(line)
                for run in p_crawled.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)


    # Footer
    doc.add_page_break()
    doc.add_paragraph("\n" * 5)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Shivam Mishra | Cyber Incident Intelligence Command Center\n").bold = True
    footer.add_run("CONFIDENTIAL SECURITY INTELLIGENCE REPORT\n")
    footer.add_run("This document is AI-generated and reviewed for forensic accuracy.")
    
    doc.save(file_path)
    return file_path


def generate_combined_pdf(report, incidents, cves, include_crawled_content=True):
    """Generates a professional combined Intelligence Report (PDF)."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"Intelligence_Report_{report.id}.pdf")
    
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=letter,
        rightMargin=40, leftMargin=40,
        topMargin=60, bottomMargin=50
    )
    styles = getSampleStyleSheet()
    
    # Custom premium styles
    title_style = ParagraphStyle(
        'CombinedTitle',
        parent=styles['Heading1'],
        fontSize=32,
        textColor=colors.HexColor("#1e1b4b"),
        alignment=TA_CENTER,
        spaceAfter=20,
        fontName='Helvetica-Bold'
    )
    
    credit_style = ParagraphStyle(
        'CreditStyle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor("#4f46e5"),
        alignment=TA_CENTER,
        spaceAfter=40,
        fontName='Helvetica-Bold'
    )
    
    section_header = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.white,
        backColor=colors.HexColor("#4f46e5"),
        fontName='Helvetica-Bold',
        spaceBefore=20,
        spaceAfter=15,
        leftIndent=0,
        borderPadding=8
    )
    
    item_title = ParagraphStyle(
        'ItemTitle',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor("#3730a3"),
        spaceBefore=10,
        spaceAfter=5
    )
    
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8
    )
    
    meta_style = ParagraphStyle(
        'MetaText',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=5
    )

    elements = []
    
    # Title Page Info
    elements.append(Spacer(1, 150))
    elements.append(Paragraph(report.report_title.upper(), title_style))
    elements.append(Paragraph("OFFICIAL SECURITY INTELLIGENCE REPORT", ParagraphStyle('SubTitle', parent=styles['Normal'], alignment=TA_CENTER, fontSize=14, textColor=colors.grey, spaceAfter=10)))
    elements.append(Paragraph("Created by Shivam Mishra", credit_style))
    date_range = f"<b>PERIOD:</b> {report.from_date.strftime('%Y-%m-%d')} TO {report.to_date.strftime('%Y-%m-%d')}"
    elements.append(Paragraph(date_range, ParagraphStyle('DateRange', parent=styles['Normal'], alignment=TA_CENTER, fontSize=11, textColor=colors.HexColor("#4f46e5"))))
    elements.append(Spacer(1, 40))
    
    # Executive Summary placeholder or report stats
    elements.append(Paragraph("EXECUTIVE SUMMARY", section_header))
    summary_data = [
        [f"Total Incidents Identified:", f"{len(incidents)}"],
        [f"Total Vulnerabilities (CVE):", f"{len(cves)}"],
        [f"Report Generation Date:", f"{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}"],
        [f"Classification:", f"CONFIDENTIAL / INTERNAL USE ONLY"]
    ]
    st = Table(summary_data, colWidths=[200, 300])
    st.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f9fafb")),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(st)
    elements.append(PageBreak())
    
    # Section 1: Incident Intelligence
    if incidents:
        elements.append(Paragraph("SECTION 1: CYBER INCIDENT INTELLIGENCE", section_header))
        for idx, inc in enumerate(incidents):
            elements.append(Paragraph(f"1.{idx+1} {inc.title}", item_title))
            meta = f"<b>Date:</b> {inc.happened_at.strftime('%Y-%m-%d') if inc.happened_at else 'N/A'} | <b>Severity:</b> {inc.severity} | <b>Source:</b> {inc.source}"
            elements.append(Paragraph(meta, meta_style))
            if inc.link:
                elements.append(Paragraph(f"<b>Source URL:</b> {inc.link}", meta_style))
            elements.append(Paragraph(inc.ai_summary or inc.description or "No detailed analysis.", body_style))
            if include_crawled_content and inc.crawled_content:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("<b>Source Article Details (Crawled Content):</b>", body_style))
                # Truncate very long crawled content to keep PDF manageable
                crawled_text = inc.crawled_content[:3000] + ('...' if len(inc.crawled_content) > 3000 else '')
                elements.append(Paragraph(crawled_text, body_style))
            elements.append(Spacer(1, 10))
            if idx < len(incidents) - 1:
                elements.append(Spacer(1, 5))

    # Section 2: Vulnerability Landscape
    if cves:
        if incidents: elements.append(PageBreak())
        elements.append(Paragraph("SECTION 2: VULNERABILITY LANDSCAPE (CVE/NVD)", section_header))
        for idx, cve in enumerate(cves):
            elements.append(Paragraph(f"2.{idx+1} {cve.cve_id} - {cve.company_name or 'Global Vulnerability'}", item_title))
            meta = f"<b>Published:</b> {cve.published_date.strftime('%Y-%m-%d') if cve.published_date else 'N/A'} | <b>Severity:</b> {cve.severity} | <b>Score:</b> {cve.cvss_score or 'N/A'}"
            elements.append(Paragraph(meta, meta_style))
            elements.append(Paragraph("<b>Raw NVD Description:</b>", body_style))
            elements.append(Paragraph(cve.description or "No description available.", body_style))
            if cve.ai_summary:
                elements.append(Paragraph("<b>AI Intelligence Summary:</b>", body_style))
                elements.append(Paragraph(cve.ai_summary, body_style))
            elements.append(Spacer(1, 10))

    doc.build(elements, onFirstPage=_draw_report_header_footer, onLaterPages=_draw_report_header_footer)
    return file_path


def generate_combined_docx(report, incidents, cves, include_crawled_content=True):
    """Generates a professional combined Intelligence Report (Word)."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"Intelligence_Report_{report.id}.docx")
    
    doc = Document()
    
    # Title Page
    title = doc.add_heading('', 0)
    run = title.add_run(report.report_title.upper())
    run.font.size = Pt(36)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("OFFICIAL SECURITY INTELLIGENCE REPORT")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    
    p_credit = doc.add_paragraph()
    p_credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cred = p_credit.add_run("Created by Shivam Mishra")
    run_cred.bold = True
    run_cred.font.size = Pt(14)
    run_cred.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PERIOD: {report.from_date.strftime('%Y-%m-%d')} TO {report.to_date.strftime('%Y-%m-%d')}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    
    doc.add_paragraph("\n" * 5)
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    doc.add_paragraph(f"Report Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total Incidents: {len(incidents)}")
    doc.add_paragraph(f"Total CVEs: {len(cves)}")
    
    # Incidents
    if incidents:
        doc.add_page_break()
        doc.add_heading('SECTION 1: CYBER INCIDENT INTELLIGENCE', level=1)
        for idx, inc in enumerate(incidents):
            doc.add_heading(f"1.{idx+1} {inc.title}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"Date: {inc.happened_at.strftime('%Y-%m-%d') if inc.happened_at else 'N/A'} | Severity: {inc.severity} | Source: {inc.source}").italic = True
            if inc.link:
                p_link = doc.add_paragraph()
                p_link.add_run("Source URL: ").bold = True
                p_link.add_run(inc.link)
            doc.add_paragraph(inc.ai_summary or inc.description or "No detailed analysis.")
            if include_crawled_content and inc.crawled_content:
                p_crawl = doc.add_paragraph()
                p_crawl.add_run("Source Article Details (Crawled Content): ").bold = True
                crawled_text = inc.crawled_content[:3000] + ('...' if len(inc.crawled_content) > 3000 else '')
                p_crawl.add_run(crawled_text)
            
    # CVEs
    if cves:
        doc.add_page_break()
        doc.add_heading('SECTION 2: VULNERABILITY LANDSCAPE (CVE/NVD)', level=1)
        for idx, cve in enumerate(cves):
            doc.add_heading(f"2.{idx+1} {cve.cve_id} - {cve.company_name or 'Global Vulnerability'}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"Published: {cve.published_date.strftime('%Y-%m-%d') if cve.published_date else 'N/A'} | Severity: {cve.severity} | Score: {cve.cvss_score or 'N/A'}").italic = True
            p2 = doc.add_paragraph()
            p2.add_run("Raw NVD Description: ").bold = True
            p2.add_run(cve.description or "No description available.")
            if cve.ai_summary:
                p3 = doc.add_paragraph()
                p3.add_run("AI Intelligence Summary: ").bold = True
                p3.add_run(cve.ai_summary)
                
    doc.save(file_path)
    return file_path


def generate_combined_excel(report, incidents, cves, include_crawled_content=True):
    """Generates a structured multi-sheet Excel report for combined data."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"Intelligence_Report_{report.id}.xlsx")
    
    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        # Sheet 1: Report Metadata
        metadata = pd.DataFrame([{
            "Report Title": report.report_title,
            "Period From": report.from_date.strftime('%Y-%m-%d'),
            "Period To": report.to_date.strftime('%Y-%m-%d'),
            "Generated At": pd.Timestamp.now().strftime('%Y-%m-%d %H:%M'),
            "Author": "Shivam Mishra",
            "Classification": "Confidential Security Intelligence"
        }])
        metadata.to_excel(writer, sheet_name='Summary & Meta', index=False)
        if incidents:
            inc_data = [{
                "Date": i.happened_at.strftime("%Y-%m-%d") if i.happened_at else "N/A",
                "Severity": i.severity,
                "Title": i.title,
                "Source": i.source,
                "Source URL": i.link,
                "Description": i.description,
                "AI Summary": i.ai_summary,
                "Attack Type": i.attack_type,
                "Crawled Content": i.crawled_content if include_crawled_content else ""
            } for i in incidents]
            pd.DataFrame(inc_data).to_excel(writer, sheet_name="Incidents", index=False)
            
        if cves:
            cve_data = [{
                "CVE ID": c.cve_id,
                "Severity": c.severity,
                "Score": c.cvss_score,
                "Company": c.company_name,
                "Product": c.product_name,
                "Published": c.published_date.strftime("%Y-%m-%d") if c.published_date else "N/A",
                "Description": c.description,
                "AI Summary": c.ai_summary
            } for c in cves]
            pd.DataFrame(cve_data).to_excel(writer, sheet_name="Vulnerabilities", index=False)
            
        # Summary Sheet
        summary_data = [
            {"Parameter": "Report Title", "Value": report.report_title},
            {"Parameter": "From Date", "Value": report.from_date.strftime("%Y-%m-%d")},
            {"Parameter": "To Date", "Value": report.to_date.strftime("%Y-%m-%d")},
            {"Parameter": "Generated On", "Value": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")},
            {"Parameter": "Total Incidents", "Value": len(incidents)},
            {"Parameter": "Total CVEs", "Value": len(cves)}
        ]
        pd.DataFrame(summary_data).to_excel(writer, sheet_name="Summary", index=False)

    return file_path

def _draw_mitre_report_footer(canvas, doc):
    canvas.saveState()
    # Header Accent line
    canvas.setFillColor(colors.HexColor("#3b82f6")) # Blue
    canvas.rect(0, letter[1] - 4, letter[0], 4, fill=1, stroke=0)
    
    # Footer
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor("#94a3b8"))
    footer_text = "Generated by Shivam AI | This report was extracted directly from MITRE CVE Records"
    canvas.drawString(40, 30, footer_text)
    canvas.line(40, 45, letter[0]-40, 45)
    
    # Page Number
    page_num = canvas.getPageNumber()
    canvas.drawRightString(letter[0]-40, 30, f"Page {page_num}")
    canvas.restoreState()

def generate_single_cve_pdf(cve):
    """
    Generates a professional, high-fidelity PDF report for a single CVE vulnerability,
    matching the MITRE CVE Records UI layout.
    """
    import dateutil.parser
    import dataclasses
    import json
    
    mitre_report = None
    try:
        if cve.raw_data and "cveMetadata" in cve.raw_data:
            from Cve_parser import CVEParser
            parser = CVEParser(json.dumps(cve.raw_data))
            mitre_report = parser.generate_report()
            if 'affected' in mitre_report and 'products' in mitre_report['affected']:
                mitre_report['affected']['products'] = [dataclasses.asdict(p) for p in mitre_report['affected']['products']]
    except Exception as e:
        print("MITRE parsing error:", e)
        mitre_report = None

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"CVE_Report_{cve.cve_id.replace('-', '_')}.pdf")

    from reportlab.platypus import HRFlowable

    doc = SimpleDocTemplate(
        file_path, 
        pagesize=letter,
        rightMargin=40, leftMargin=40,
        topMargin=50, bottomMargin=60
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Modern Styles
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor("#0f172a"), spaceAfter=2, fontName="Helvetica-Bold", leading=28)
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor("#64748b"), spaceAfter=15, fontName="Helvetica")
    
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor("#1e293b"), spaceBefore=18, spaceAfter=10, fontName="Helvetica-Bold")
    h3_style = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=12, textColor=colors.HexColor("#334155"), spaceBefore=12, spaceAfter=6, fontName="Helvetica-Bold")
    
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10.5, leading=16, textColor=colors.HexColor("#334155"))
    body_bold = ParagraphStyle('BodyBold', parent=body_style, fontName="Helvetica-Bold")
    
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor("#64748b"), textTransform='uppercase', spaceAfter=2, fontName="Helvetica-Bold")
    
    mono_style = ParagraphStyle('Mono', parent=styles['Normal'], fontSize=9, fontName="Courier", textColor=colors.HexColor("#0f172a"), wordWrap='CJK')
    mono_blue = ParagraphStyle('MonoBlue', parent=mono_style, textColor=colors.HexColor("#2563eb"))
    
    link_style = ParagraphStyle('Link', parent=styles['Normal'], fontSize=9.5, fontName="Courier", textColor=colors.HexColor("#3b82f6"), wordWrap='CJK')
    
    elements = []
    
    # Extract data with fallbacks
    title = mitre_report["vulnerability"]["title"] if mitre_report else "No Title Available"
    if not title or title == "N/A": title = "No Title Available"
    
    vendor_product = f"{cve.company_name or 'Unknown'} / {cve.product_name or 'Unknown'}"
    if mitre_report and mitre_report.get("affected", {}).get("products"):
        p = mitre_report["affected"]["products"][0]
        vendor_product = f"{p.get('vendor', 'Unknown')} {p.get('product', '')}".strip()
    
    state = "PUBLISHED"
    
    pub_date = str(cve.published_date).split(' ')[0] if cve.published_date else "Unknown"
    upd_date = str(cve.last_modified_date).split(' ')[0] if cve.last_modified_date else "Unknown"
    
    desc = cve.description or "No description available."
    cwe_id = "N/A"
    cwe_desc = ""
    if mitre_report and mitre_report["vulnerability"]["cwes"]:
        cwe_id = mitre_report["vulnerability"]["cwes"][0]["id"]
        cwe_desc = mitre_report["vulnerability"]["cwes"][0]["description"]
    
    cvss_score = cve.cvss_score or "N/A"
    sev = (cve.severity or "UNKNOWN").upper()
    
    # Severity Badge Color
    sev_bg = "#f59e0b"
    if sev == "CRITICAL": sev_bg = "#ef4444"
    elif sev == "HIGH": sev_bg = "#f97316"
    elif sev == "MEDIUM": sev_bg = "#eab308"
    elif sev == "LOW": sev_bg = "#22c55e"
    
    attack_vector = "Unknown"
    cvss_vector = cve.raw_data.get("metrics", {}).get("cvssMetricV31", [{}])[0].get("cvssData", {}).get("vectorString", "Unknown") if cve.raw_data else "Unknown"
    if cvss_vector == "Unknown" and mitre_report:
        cvss_vector = mitre_report["scoring"]["cvss_vector"] or "Unknown"
        
    if cvss_vector != "Unknown":
        if "AV:N" in cvss_vector: attack_vector = "Network"
        elif "AV:L" in cvss_vector: attack_vector = "Local"
        elif "AV:A" in cvss_vector: attack_vector = "Adjacent"
        elif "AV:P" in cvss_vector: attack_vector = "Physical"
        
    prod_count = len(cve.affected_products) if cve.affected_products else 0
    if mitre_report:
        prod_count = mitre_report["affected"].get("product_count", prod_count)

    # Header Section
    elements.append(Paragraph("VULNERABILITY REPORT", label_style))
    title_p = Paragraph(title, title_style)
    sub_p = Paragraph(f"<b>{cve.cve_id}</b> &bull; {vendor_product} &bull; {state}", subtitle_style)
    
    badge_data = [[
        Paragraph(f"<font size=10 color='white'><b>CVSS 3.1</b></font><br/><font size=24 color='white'><b>{cvss_score}</b></font><br/><font size=10 color='white'><b>{sev}</b></font>", ParagraphStyle('b', alignment=TA_CENTER))
    ]]
    badge_table = Table(badge_data, colWidths=[80], rowHeights=[70])
    badge_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,0), colors.HexColor(sev_bg)),
        ('ALIGN', (0,0), (0,0), 'CENTER'),
        ('VALIGN', (0,0), (0,0), 'MIDDLE'),
        ('ROUNDEDCORNERS', [8, 8, 8, 8]),
    ]))

    header_layout = Table([[ [title_p, sub_p], badge_table ]], colWidths=[400, 100])
    header_layout.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (1,0), (1,0), 'RIGHT'),
    ]))
    elements.append(header_layout)
    
    # Custom Divider
    elements.append(Spacer(1, 10))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=0, spaceAfter=20))
    
    # Description
    elements.append(Paragraph("Vulnerability Description", h2_style))
    elements.append(Paragraph(desc, body_style))
    elements.append(Spacer(1, 20))
    
    # Grid info
    grid_data = [
        [Paragraph("CWE ID", label_style), Paragraph("PUBLISHED", label_style)],
        [Paragraph(f"<b>{cwe_id}</b><br/><font size=9 color='#64748b'>{cwe_desc}</font>", body_style), Paragraph(f"<b>{pub_date}</b><br/><font size=9 color='#64748b'>Updated: {upd_date}</font>", body_style)],
        [Paragraph("AFFECTED PRODUCTS", label_style), Paragraph("ATTACK VECTOR", label_style)],
        [Paragraph(f"<b>{prod_count} Products</b><br/><font size=9 color='#64748b'>Multiple Vendors</font>", body_style), Paragraph(f"<b>{attack_vector}</b><br/><font size=9 color='#64748b'>CVSS Vector Context</font>", body_style)]
    ]
    grid_table = Table(grid_data, colWidths=[250, 250])
    grid_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
        ('LINEABOVE', (0,2), (-1,2), 1, colors.HexColor("#e2e8f0")),
        ('LINEBEFORE', (1,0), (1,-1), 1, colors.HexColor("#e2e8f0")),
        ('PADDING', (0, 0), (-1, -1), 12),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,0), 2),
        ('BOTTOMPADDING', (0,2), (-1,2), 2),
    ]))
    elements.append(grid_table)
    elements.append(Spacer(1, 20))
    
    # CVSS Vector Block
    if cvss_vector and cvss_vector != "Unknown":
        vec_data = [[Paragraph(f"<font color='#64748b'><b>CVSS VECTOR</b></font>", ParagraphStyle('v1', fontSize=8)), Paragraph(cvss_vector, mono_blue)]]
        vec_table = Table(vec_data, colWidths=[90, 410])
        vec_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f1f5f9")),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#cbd5e1")),
            ('PADDING', (0,0), (-1,-1), 8),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        elements.append(vec_table)
        elements.append(Spacer(1, 20))

    # Affected Products List
    elements.append(Paragraph(f"Affected Systems ({prod_count} total)", h2_style))
    
    if mitre_report and mitre_report.get("affected", {}).get("products"):
        for prod in mitre_report["affected"]["products"]:
            elements.append(Paragraph(f"<b>{prod.get('vendor','')} {prod.get('product','')}</b>", body_style))
            platforms = ", ".join(prod.get("platforms", [])) if prod.get("platforms") else "No platforms specified"
            elements.append(Paragraph(f"<font color='#64748b'>Platforms: {platforms}</font>", body_style))
            elements.append(Spacer(1, 4))
            
            if prod.get("versions"):
                ver_table_data = []
                for v in prod["versions"]:
                    status_col = "#22c55e" if v.get("status") == "unaffected" else "#ef4444"
                    ver_str = f"&ge; {v.get('version','*')}  &rarr;  &lt; {v.get('less_than','*')}"
                    ver_table_data.append([
                        Paragraph(ver_str, mono_style),
                        Paragraph(f"<font color='{status_col}'><b>{v.get('status','affected').upper()}</b></font>", ParagraphStyle('status', fontName='Helvetica-Bold', fontSize=8, alignment=TA_CENTER))
                    ])
                vt = Table(ver_table_data, colWidths=[380, 80])
                vt.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),
                    ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
                    ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
                    ('PADDING', (0,0), (-1,-1), 6),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ]))
                elements.append(vt)
            else:
                elements.append(Paragraph("All versions affected.", body_style))
            elements.append(Spacer(1, 15))
    else:
        # Fallback to CPEs
        if cve.affected_products:
            cpe_data = [[Paragraph(cpe, mono_style)] for cpe in cve.affected_products[:15]]
            if cpe_data:
                cpe_t = Table(cpe_data, colWidths=[500])
                cpe_t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),
                    ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
                    ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
                    ('PADDING', (0,0), (-1,-1), 6),
                ]))
                elements.append(cpe_t)
        else:
            elements.append(Paragraph("No specific products listed.", body_style))
            
    elements.append(Spacer(1, 15))
    
    # References
    if cve.references:
        elements.append(Paragraph("Intel & Advisory References", h2_style))
        for ref in cve.references[:10]:
            elements.append(Paragraph(f"&bull; <a href='{ref}'>{ref}</a>", link_style))
            elements.append(Spacer(1, 6))
            
    # Build doc
    doc.build(elements, onFirstPage=_draw_mitre_report_footer, onLaterPages=_draw_mitre_report_footer)
    return file_path

import datetime

def generate_automation_summary_pdf(db, timeframe_hours=24):
    """
    Generates a PDF summary of the automation cycle's Impact Radar results.
    """
    import tempfile
    import os
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    
    import models

    cutoff_time = datetime.datetime.utcnow() - datetime.timedelta(hours=timeframe_hours)
    
    logs = db.query(models.AutomationAuditLog).filter(
        models.AutomationAuditLog.created_at >= cutoff_time
    ).order_by(models.AutomationAuditLog.created_at.desc()).all()
    
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"Automation_Summary_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
    
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=letter,
        rightMargin=40, leftMargin=40,
        topMargin=50, bottomMargin=50
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('MainTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor("#0a0e17"))
    header_style = ParagraphStyle('SectionHeader', parent=styles['Heading2'], fontSize=12, textColor=colors.HexColor("#1e1b4b"), spaceBefore=15, spaceAfter=8)
    body_style = ParagraphStyle('ReportBody', parent=styles['Normal'], fontSize=9, leading=12)
    
    elements = []
    
    # Title
    elements.append(Paragraph("AUTOMATION & IMPACT RADAR EXECUTION REPORT", title_style))
    elements.append(Paragraph(f"Reporting Period: Last {timeframe_hours} Hours", body_style))
    elements.append(Paragraph(f"Generated On: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} UTC", body_style))
    elements.append(Spacer(1, 20))
    
    # Summary Metrics
    total_scanned = len(logs)
    matched_logs = [log for log in logs if log.impact_score >= 60]
    rejected_logs = [log for log in logs if log.impact_score < 60]
    
    summary_data = [
        [Paragraph("<b>Metric</b>", body_style), Paragraph("<b>Value</b>", body_style)],
        [Paragraph("Total Items Scanned", body_style), Paragraph(str(total_scanned), body_style)],
        [Paragraph("Impact Radar Matches (>= 60)", body_style), Paragraph(f"<font color='red'><b>{len(matched_logs)}</b></font>", body_style)],
        [Paragraph("Impact Radar Rejections (< 60)", body_style), Paragraph(f"<font color='green'><b>{len(rejected_logs)}</b></font>", body_style)]
    ]
    t_summary = Table(summary_data, colWidths=[200, 100])
    t_summary.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f3f4f6")),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t_summary)
    elements.append(Spacer(1, 20))
    
    # Matches Table
    elements.append(Paragraph("IMPACT RADAR MATCHES (High Risk / Pushed to Jira)", header_style))
    if matched_logs:
        match_data = [[Paragraph("<b>ID</b>", body_style), Paragraph("<b>Type</b>", body_style), Paragraph("<b>Target/Vendor</b>", body_style), Paragraph("<b>Score</b>", body_style), Paragraph("<b>Match Details</b>", body_style)]]
        for log in matched_logs[:50]: # Limit to 50 for brevity
            reason = "N/A"
            if log.entity_type.lower() == 'cve':
                cve_obj = db.query(models.CVE).filter(models.CVE.cve_id == log.entity_id).first()
                if cve_obj and cve_obj.company_impact_reason: reason = cve_obj.company_impact_reason
            elif log.entity_type.lower() == 'incident':
                try:
                    inc_obj = db.query(models.Incident).filter(models.Incident.id == int(log.entity_id)).first()
                    if inc_obj and inc_obj.company_impact_reason: reason = inc_obj.company_impact_reason
                except:
                    pass
                    
            match_data.append([
                Paragraph(log.entity_id, body_style),
                Paragraph(log.entity_type.upper(), body_style),
                Paragraph(log.entity_title or "Unknown", body_style),
                Paragraph(f"<font color='red'><b>{log.impact_score}</b></font>", body_style),
                Paragraph(reason, body_style)
            ])
        t_matches = Table(match_data, colWidths=[80, 50, 150, 40, 200])
        t_matches.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#fee2e2")),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(t_matches)
    else:
        elements.append(Paragraph("No items matched the Impact Radar criteria in this timeframe.", body_style))
        
    elements.append(Spacer(1, 20))
    
    # Rejections Table
    elements.append(Paragraph("IMPACT RADAR REJECTIONS (Low Risk / Ignored)", header_style))
    if rejected_logs:
        rej_data = [[Paragraph("<b>ID</b>", body_style), Paragraph("<b>Type</b>", body_style), Paragraph("<b>Target/Vendor</b>", body_style), Paragraph("<b>Score</b>", body_style), Paragraph("<b>Match Details</b>", body_style)]]
        for log in rejected_logs[:50]: # Limit to 50
            reason = "N/A"
            if log.entity_type.lower() == 'cve':
                cve_obj = db.query(models.CVE).filter(models.CVE.cve_id == log.entity_id).first()
                if cve_obj and cve_obj.company_impact_reason: reason = cve_obj.company_impact_reason
            elif log.entity_type.lower() == 'incident':
                try:
                    inc_obj = db.query(models.Incident).filter(models.Incident.id == int(log.entity_id)).first()
                    if inc_obj and inc_obj.company_impact_reason: reason = inc_obj.company_impact_reason
                except:
                    pass
                    
            rej_data.append([
                Paragraph(log.entity_id, body_style),
                Paragraph(log.entity_type.upper(), body_style),
                Paragraph(log.entity_title or "Unknown", body_style),
                Paragraph(str(log.impact_score), body_style),
                Paragraph(reason, body_style)
            ])
        t_rej = Table(rej_data, colWidths=[80, 50, 150, 40, 200])
        t_rej.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#dcfce7")),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(t_rej)
    else:
        elements.append(Paragraph("No items were rejected by the Impact Radar in this timeframe.", body_style))

    doc.build(elements)
    return file_path
